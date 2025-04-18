from datetime import datetime
from typing import List, Dict, Optional
from docx import Document
from docx2pdf import convert

from utils.file_work import *
from utils.json_work import *
from utils.diff_work import *
from analyze.mistral_analyze import *
from download_repo import *

WEIGHTS = {
    "Security": 3.0,
    "CodeStyle": 1.5,
    "CodeSmells": 2.0,
    "AntiPatterns": 2.5,
    "LegacyCompatibility": 1.0
}

def form_report(github_url: str,
    email: str,
    start_date: datetime,
    end_date: datetime,
    access_token: Optional[str] = None):

    # temporary_time = datetime.now().strftime("%Y-%m-%d_%H-%M")
    # temporary_dir = f'./temporary-{temporary_time}'
    time = datetime.now().strftime("%Y-%m-%d_%H-%M")
    diff_output_dir = f'diff-{time}'

    diffs = get_diffs(
        github_url=github_url,
        email=email,
        start_date=start_date,
        end_date=end_date,
        # токен гитхаба из переменных окружения
        access_token=access_token,
        output_dir=diff_output_dir
    )
    for diff in diffs:
        logger.info(f"PR #{diff['pr_number']}: {diff['title']} (Diff saved to {diff['diff_path']})")
        author = diff['author']
    logger.info(f"Найдено {len(diffs)} PR")
    
    file_count = 0
    score = 0
    history_file = f'history-{author}-{time}.txt'

    report = Document()
    report.add_heading('Отчет об оценке качества кода', 0)
    p = report.add_paragraph()
    report.add_heading(f'Автор: {author}', level=1)
    report.add_heading('Выявленные проблемы', level=1)

    for diff in diffs:
        p = report.add_paragraph()
        p.add_run('\nНОМЕР МР: ').bold = True
        p.add_run(str(diff["pr_number"]))
        p = report.add_paragraph()
        p.add_run('КОММИТ: ').bold = True
        p.add_run(str(diff["commit_sha"]))

        time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        output_dir = f'restored-{author}-{time}'

        diff_dict = analyze_diff_and_write_in_docx(diff['diff_path'], output_dir, report, history_file)
                    
        for key, value in diff_dict.items():
            file_count += key
            score += value
        
    # make_review_and_write_in_docx(report, history_file)

    if file_count != 0:
        score /= file_count * 10

        p = report.add_paragraph()
        p.add_run('\nОБЩАЯ ОЦЕНКА: ').bold = True
        p.add_run(str(score))

    report_name = f'report-{author}-{time}'
    report.save(f'{report_name}.docx')
    convert(report_name +'.docx', report_name + '.pdf')

    delete_dir(output_dir)
    delete_dir(diff_output_dir)
    delete_file(report_name +'.docx')
    delete_file(history_file)

    return report_name + '.pdf'

def analyze_diff_and_write_in_docx(diff_path: str, output_dir: str, docx: str, history_file: str) -> None:
    """This function analyzes diff file and gives feedback about it and how good the developer is
        diff_path: path to diff file
        docx: path to docx file
        history_file: path to history files
    """
    # сумма очков за каждый критерий * вес критерия
    score = 0

    # парсим диффы в файлы с кодом с сохраннием структуры папок
    parse_diff_to_code_files(diff_path, output_dir)
    diff_files = get_files_in_dir(output_dir)

    for file in diff_files:
        code = read_file(file)
        prompt = make_prompt(code)
        ai_result = mistral_analyze(prompt)
        # ЗДЕСЬ БУДЕТ СТАТИЧЕСКИЙ АНАЛИЗ НАПРИМЕР:
        # stat_result = static_analyze(code)

        # записываем результат в файл истории
        write_file(history_file, f"User: {prompt}\nAssistant: {ai_result}\n")
        # ЗДЕСЬ БУДЕТ ЗАПИСЬ РЕЗУЛЬТАТОВ СТАТ АНАЛИЗА В ФАЙЛ ИСТОРИИ НАПРИМЕР:
        # write_file(history_file, f"User: Great! Now find vulnerabilities and poor code style\nAssistant: {stat_result}\n")

        p = docx.add_paragraph()
        p.add_run(f"ФАЙЛ: {file}\n").bold = True
        write_json_to_docx_file(docx, ai_result)
        # ЗДЕСЬ БУДЕТ ЗАПИСЬ РЕЗУЛЬТАТОВ СТАТ АНАЛИЗА В docx НАПРИМЕР:
        # write_json_to_docx_file(docx, stat_result)

        for key in ai_result.keys():
            small_json = ai_result[key]
            score += small_json['score'] * WEIGHTS[key]

        # # ЗДЕСЬ БУДЕТ ДОБАВЛЕНИЕ ОЧКОВ ЗА СТАТ АНАЛИЗ В SCORE НАПРИМЕР:
        # for key in stat_result.keys():
        #     small_json = stat_result[key]
        #     score += small_json['score'] * WEIGHTS[key]


    # возвращаем словарь {кол-во файлов: общая оценка} для формирования оценки разработчика
    return {len(diff_files): score}

def make_review_and_write_in_docx(docx: str, history_file: str) -> None:
    review_prompt = """
        You are an experienced team leader and code reviewer. Your job is to give the developer feedback on the quality of their code.
        Analyze the code you've seen and what you've said about it for the following aspects:
        1. What good aspects did you notice in the code? (e.g. good structure, readability, attention to security, etc.)
        2. What improvements would you suggest? (optimization, style, architecture, security, etc.)
        3. Give 2-3 specific tips on what the developer should learn or improve to write even better
        4. If there are positive or negative aspects associated with the outdated stack, be sure to mention it.

        Make sure to use double quotes in JSON!!!! DON'T repeat the text in the example!
        Output must look like this:

        json
        {
            "Overall Review": {
                "strengths": "The code is well structured, modules are logically separated. 
                                Explicit typing is used",
                "improvements": "In places the code style is not followed - indents and line lengths. 
                                    Repetitive code should be avoided, for example, in functions X and Y",
                "recommendations": "Learn SOLID principles and start applying them in the architecture.
                Try using linters (e.g. flake8, eslint) for automatic style control.
                Understand design patterns, especially in terms of separation of concerns"
            }
        }

        Do not include 'legacy_context' and 'forced_solution' into the final response's JSON!!!
        """
    history = read_file(history_file)
    result = mistral_analyze(review_prompt, history)

    write_json_to_docx_file(docx, result)
