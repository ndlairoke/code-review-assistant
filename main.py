from download_repo import *
import os
from datetime import datetime, timezone
import logging
from file_work import *
from docx import Document
from analyze import *

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

if __name__ == "__main__":
    # Пример использования
    try:
        diffs = get_diffs(
            github_url="https://github.com/natasha/natasha",
            email="122558239+the-lenoz@users.noreply.github.com",
            start_date=datetime(2023, 1, 1, tzinfo=timezone.utc),
            end_date=datetime(2025, 4, 17, tzinfo=timezone.utc),
            # токен гитхаба из переменных окружения
            access_token=os.getenv("GITHUB_TOKEN"),
            output_dir="diffs"
        )
        for diff in diffs:
            logger.info(f"PR #{diff['pr_number']}: {diff['title']} (Diff saved to {diff['diff_path']})")
            author = diff['author']
        logger.info(f"Найдено {len(diffs)} PR")
    except Exception as e:
        logger.error(f"Ошибка: {str(e)}")

#             email="avasilev1992@gmail.com",
#             repo_url="https://github.com/AlfaInsurance/devQ_testData_PythonProject.git",

    # код Ани с обработкой диффов и все такое
    time = datetime.now().strftime("%Y-%m-%d_%H-%M")
    file_count = 0
    score = 0

    history_file = f'history-{time}.txt'
    report = Document()
    report.add_heading('Отчет об оценке качества кода', 0)
    p = report.add_paragraph()
    p.add_run('Общая информация').bold = True
    p = report.add_paragraph()
    p.add_run('Автор: ').bold = True
    p.add_run(author) # тут вставлять diff['author']

    report.add_heading('Выявленные проблемы', level=1)

    for diff in diffs:
        p = report.add_paragraph()
        p.add_run('Номер МР: ').bold = True
        p.add_run(str(diff["pr_number"]))
        p = report.add_paragraph()
        p.add_run('Коммит: ').bold = True
        p.add_run(str(diff["commit_sha"]))

        diff_dict = analyze_diff_and_write_in_docx(diff['diff_path'], report, history_file)
                
        for key, value in diff_dict.items():
            file_count += key
            score += value
    
    make_review_and_write_in_docx(report, history_file)

    if file_count != 0:
        score /= file_count

        p = report.add_paragraph()
        p.add_run('Общая оценка: ').bold = True
        p.add_run(str(score))

    report_name = f'report-{time}'
    report.save(report_name + '.docx')
    convert(report_name + '.docx', report_name + '.pdf')

    delete_file(history_file)
    delete_file(report_name + '.docx')

    print(f'Отчет сгенерирован в файле {report_name + '.pdf'}')
